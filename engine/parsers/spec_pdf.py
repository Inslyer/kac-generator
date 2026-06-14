"""Парсер спецификации → список SpecPosition. Принимает PDF, DOCX, XLSX/XLS.

Форматы:
  • PDF с текстовым слоем → текст напрямую (PyMuPDF).
  • PDF-скан без текста → локальный OCR (Tesseract/Apple Vision); если OCR нет, а провайдер
    anthropic — фолбэк на Claude vision (картинки).
  • DOCX → текст абзацев и таблиц (python-docx).
  • XLSX/XLS → строки листов как текст (openpyxl).

Структуризация выполняется LLM (Claude или DeepSeek). Без ключа — бросает LLMUnavailable,
UI предлагает ручной ввод. Завод-изготовитель из CUSTOM_MAKERS → авто-пометка уникального.
"""
from __future__ import annotations

import io
from pathlib import Path

import fitz

from ..config import CACHE_DIR, settings
from ..llm import LLMUnavailable, complete_json
from ..models import SpecPosition

SYSTEM = (
    "Ты извлекаешь позиции из спецификации оборудования (рабочая документация). "
    "Верни JSON-массив объектов с полями: number (int), name (string — наименование и "
    "техническая характеристика), type_mark (string — тип/марка/обозначение, если есть), "
    "unit (string — ед. изм., напр. шт/м/компл), qty (number — количество), "
    "maker (string — завод-изготовитель, если указан). "
    "Подпозиции «в составе» объединяй в родительскую позицию. "
    "Не выдумывай позиции. Сохраняй порядок. Только JSON."
)


def extract_raw_text(pdf_path: str | Path) -> tuple[str, bool]:
    """Возвращает (текст, is_scanned). Для сканов текст пустой → нужен OCR/vision."""
    doc = fitz.open(pdf_path)
    parts = [doc[i].get_text() for i in range(doc.page_count)]
    text = "\n".join(parts)
    is_scanned = len(text.strip()) < 40 * doc.page_count  # эвристика
    doc.close()
    return text, is_scanned


def text_from_docx(path: str | Path) -> str:
    """Текст из DOCX: абзацы + ячейки таблиц (таблицы — построчно, через табуляцию)."""
    import docx

    doc = docx.Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def text_from_xlsx(path: str | Path, max_rows: int = 4000) -> str:
    """Текст из XLSX/XLS: строки всех листов через табуляцию (пустые строки пропускаются)."""
    import openpyxl

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Лист: {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                break
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append("\t".join(cells))
    wb.close()
    return "\n".join(parts)


def render_pages(pdf_path: str | Path, dpi: int = 170, max_pages: int = 12) -> list[str]:
    """Рендерит страницы PDF в PNG (для распознавания сканов через Claude vision)."""
    doc = fitz.open(pdf_path)
    paths = []
    for i in range(min(max_pages, doc.page_count)):
        out = CACHE_DIR / f"spec_{Path(pdf_path).stem}_{i+1}.png"
        doc[i].get_pixmap(dpi=dpi).save(out)
        paths.append(str(out))
    doc.close()
    return paths


def _ocr_backend() -> str:
    """Доступный локальный OCR: 'vision' (Apple, macOS) → 'tesseract' → '' (нет)."""
    try:
        import ocrmac  # noqa: F401
        return "vision"
    except Exception:
        pass
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return "tesseract"
    except Exception:
        return ""


def ocr_available() -> bool:
    """Доступно ли локальное распознавание сканов (Apple Vision или Tesseract)."""
    return bool(_ocr_backend())


def ocr_pages(pdf_path: str | Path, dpi: int = 300, max_pages: int = 12) -> str:
    """Локальное распознавание скана в текст. Высокий DPI для таблиц.

    На macOS — Apple Vision (ocrmac, без системных зависимостей); иначе — Tesseract.
    """
    from PIL import Image

    backend = _ocr_backend()
    doc = fitz.open(pdf_path)
    parts: list[str] = []
    try:
        for i in range(min(max_pages, doc.page_count)):
            pix = doc[i].get_pixmap(dpi=dpi)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            if backend == "vision":
                from ocrmac import ocrmac as _ocrmac
                # Apple Vision: список (текст, уверенность, bbox) сверху вниз
                ann = _ocrmac.OCR(img, language_preference=["ru-RU", "en-US"]).recognize()
                parts.append("\n".join(a[0] for a in ann))
            elif backend == "tesseract":
                import pytesseract
                parts.append(pytesseract.image_to_string(img, lang="rus+eng"))
    finally:
        doc.close()
    return "\n".join(parts)


def _structure_positions(text: str = "", images: list[str] | None = None) -> list[dict]:
    """Зовёт LLM (текст или vision) и возвращает сырой массив позиций."""
    # лимит большой: xlsx/docx-спецификации бывают объёмными, DeepSeek/Claude держат контекст
    prompt = (f"Спецификация (текст):\n\n{text[:120000]}" if text
              else "Распознай таблицу-спецификацию со страниц-изображений.")
    # быстрая модель: на сканах (vision) это в разы быстрее, качество таблиц достаточное
    return complete_json(prompt, system=SYSTEM, smart=False, max_tokens=8192, images=images)


def _raw_positions_from_pdf(path: str | Path) -> list[dict]:
    text, is_scanned = extract_raw_text(path)
    if not is_scanned:
        return _structure_positions(text=text)
    # Скан: сперва локальный OCR (работает с любым LLM, в т.ч. DeepSeek без vision).
    ocr_text = ocr_pages(path) if ocr_available() else ""
    if ocr_text.strip():
        return _structure_positions(text=ocr_text)
    if settings.llm_provider.lower() == "anthropic":
        return _structure_positions(images=render_pages(path))  # фолбэк на vision
    raise LLMUnavailable(
        "Скан без текстового слоя, локальный OCR недоступен, а провайдер "
        f"{settings.llm_provider} не умеет распознавать изображения. Установите Tesseract "
        "(brew install tesseract tesseract-lang) или используйте текстовый PDF/DOCX/XLSX.")


def _raw_positions(path: str | Path) -> list[dict]:
    """Извлекает сырой массив позиций из файла по его расширению (PDF/DOCX/XLSX)."""
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return _raw_positions_from_pdf(path)
    if suffix == ".docx":
        return _structure_positions(text=text_from_docx(path))
    if suffix in (".xlsx", ".xlsm", ".xls"):
        return _structure_positions(text=text_from_xlsx(path))
    raise LLMUnavailable(
        f"Неподдерживаемый формат спецификации: {suffix or 'без расширения'}. "
        "Принимаются PDF, DOCX, XLSX.")


def parse_spec(
    path: str | Path,
    discipline: str = "ЭМ",
    unique_names: set[str] | None = None,
) -> list[SpecPosition]:
    """Парсит спецификацию (PDF/DOCX/XLSX) в список позиций.

    unique_names — подстроки наименований, помечаемые как уникальное оборудование.
    Дополнительно: завод-изготовитель из CUSTOM_MAKERS → авто-пометка уникальным.
    """
    raw = _raw_positions(path)

    unique_names = unique_names or set()
    positions: list[SpecPosition] = []
    for item in raw:
        name = str(item.get("name", "")).strip()
        if not name:
            continue
        maker = str(item.get("maker", "")).strip()
        is_unique = (
            any(u.lower() in name.lower() for u in unique_names)
            or any(m in maker.lower() for m in settings.custom_makers_list)
        )
        positions.append(SpecPosition(
            number=int(item.get("number", len(positions) + 1)),
            name=name,
            type_mark=str(item.get("type_mark", "")),
            unit=str(item.get("unit", "шт")) or "шт",
            qty=float(item.get("qty", 0) or 0),
            is_unique=is_unique,
            discipline=discipline,
        ))
    return positions
